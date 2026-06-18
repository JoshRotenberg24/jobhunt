#!/usr/bin/env python3
"""
render_cover_letter.py — build a professional, ATS-safe one-page cover letter.

Outputs <base>.pdf (ReportLab) and <base>.docx (python-docx), matching the
resume letterhead. Warns if it exceeds one page. Supports the same two looks as
render_resume.py — pass the style as the 2nd CLI arg or a "style" key in the JSON
(default "modern"); use the SAME style as the resume so the letterhead matches.

  - "modern"  sans-serif (Helvetica/Calibri) + navy accents.
  - "classic" serif (Times / Times New Roman), black, no color.

JSON schema:
{
  "style": "modern|classic",           # optional; default modern (CLI arg overrides)
  "name": "Joshua Rotenberg",
  "contact": {"location":"","phone":"","email":"","linkedin":""},
  "date": "June 17, 2026",
  "company": "Acme Inc.",
  "role": "Revenue Operations Manager",
  "recipient": "Hiring Team",          # optional
  "salutation": "Dear Hiring Team,",   # optional; derived from recipient if omitted
  "body": ["paragraph one", "paragraph two", "paragraph three"],
  "closing": "Sincerely,"              # optional
}

Usage:  python3 build/render_cover_letter.py <cover-letter.json> [modern|classic]
"""
import json, os, sys

INK = (0x1a, 0x1a, 0x1a); MUTED = (0x55, 0x55, 0x55)
MARGIN_TB_IN = 0.8; MARGIN_LR_IN = 0.9; BODY_PT = 10.8

THEMES = {
    "modern": {
        "pdf_regular": "Helvetica", "pdf_bold": "Helvetica-Bold",
        "docx_font": "Calibri", "accent": (0x16, 0x3a, 0x5f),
    },
    "classic": {
        "pdf_regular": "Times-Roman", "pdf_bold": "Times-Bold",
        "docx_font": "Times New Roman", "accent": INK,
    },
}


def resolve_theme(data, argv_style=None):
    name = (argv_style or data.get("style") or "modern").lower()
    if name not in THEMES:
        name = "modern"
    return name, THEMES[name]


_SUBS = {"→": "->", "←": "<-", "⇒": "=>", "↦": "->",
         "−": "-", "‒": "-", "―": "-", "‐": "-", "‑": "-",
         "·": "-", "▪": "-", "●": "-", "◦": "-", "‣": "-",
         " ": " ", " ": " ", " ": " ", " ": " "}


def _san(s):
    """Down-map characters outside WinAnsi (CP1252) coverage to ASCII so pasted
    glyphs don't render as boxes. Anything still uncovered falls back to '?'."""
    s = str(s)
    for k, v in _SUBS.items():
        s = s.replace(k, v)
    return s.encode("cp1252", "replace").decode("cp1252")


def _esc(s):
    """Sanitize to WinAnsi, then escape for ReportLab's mini-XML parser
    (e.g. company 'Procter & Gamble')."""
    return _san(s).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def build_pdf(data, path, theme):
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.units import inch
    from reportlab.lib.colors import Color
    from reportlab.lib.enums import TA_CENTER, TA_LEFT
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.platypus import (BaseDocTemplate, PageTemplate, Frame, Paragraph,
                                    Spacer, HRFlowable)
    F_REG, F_BOLD = theme["pdf_regular"], theme["pdf_bold"]
    accent = Color(*[c/255 for c in theme["accent"]]); ink = Color(*[c/255 for c in INK])
    muted = Color(*[c/255 for c in MUTED])
    pw, ph = letter; ml = MARGIN_LR_IN*inch; mt = MARGIN_TB_IN*inch
    usable = pw-2*ml; frameH = ph-2*mt

    name_st = ParagraphStyle("n", fontName=F_BOLD, fontSize=18, textColor=accent,
                             alignment=TA_CENTER, leading=21, spaceAfter=1)
    contact_st = ParagraphStyle("c", fontName=F_REG, fontSize=9, textColor=muted,
                                alignment=TA_CENTER, leading=12, spaceAfter=2)
    body = ParagraphStyle("b", fontName=F_REG, fontSize=BODY_PT, textColor=ink,
                          leading=BODY_PT+4, spaceAfter=8, alignment=TA_LEFT)
    tight = ParagraphStyle("t", fontName=F_REG, fontSize=BODY_PT, textColor=ink,
                           leading=BODY_PT+3, spaceAfter=2)

    c = data.get("contact", {})
    bits = [c.get("location"), c.get("phone"), c.get("email"), c.get("linkedin")]
    sal = data.get("salutation") or ("Dear %s," % data.get("recipient", "Hiring Team"))

    story = [Paragraph(_esc(data["name"]), name_st),
             Paragraph("&nbsp;&nbsp;&bull;&nbsp;&nbsp;".join(_esc(b) for b in bits if b), contact_st),
             HRFlowable(width="100%", thickness=0.7, color=accent, spaceBefore=2, spaceAfter=10)]
    if data.get("date"):
        story.append(Paragraph(_esc(data["date"]), tight)); story.append(Spacer(1, 6))
    if data.get("company"):
        story.append(Paragraph("<b>%s</b>" % _esc(data["company"]), tight))
    if data.get("role"):
        story.append(Paragraph("Re: %s" % _esc(data["role"]), tight))
    story.append(Spacer(1, 8))
    story.append(Paragraph(_esc(sal), body))
    for para in data.get("body", []):
        story.append(Paragraph(_esc(para), body))
    story.append(Spacer(1, 4))
    story.append(Paragraph(_esc(data.get("closing", "Sincerely,")), tight))
    story.append(Spacer(1, 14))
    story.append(Paragraph(_esc(data["name"]), tight))

    class _Doc(BaseDocTemplate):
        _last_page = 1
        def afterFlowable(self, f): self._last_page = self.canv.getPageNumber()
    doc = _Doc(path, pagesize=letter, leftMargin=ml, rightMargin=ml,
               topMargin=mt, bottomMargin=mt)
    doc.addPageTemplates([PageTemplate(id="m", frames=[Frame(ml, mt, usable, frameH,
                          leftPadding=0, rightPadding=0, topPadding=0, bottomPadding=0)])])
    doc.build(story)
    return doc._last_page


def build_docx(data, path, theme):
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    DFONT = theme["docx_font"]
    ink = RGBColor(*INK); accent = RGBColor(*theme["accent"]); muted = RGBColor(*MUTED)
    doc = Document()
    st = doc.styles["Normal"]; st.font.name = DFONT; st.font.size = Pt(BODY_PT)
    st.font.color.rgb = ink
    rf = st.element.get_or_add_rPr().get_or_add_rFonts()
    for a in ("w:ascii", "w:hAnsi", "w:cs"): rf.set(qn(a), DFONT)
    st.paragraph_format.line_spacing = 1.08
    s = doc.sections[0]
    s.top_margin = s.bottom_margin = Inches(MARGIN_TB_IN)
    s.left_margin = s.right_margin = Inches(MARGIN_LR_IN)

    def para(text, *, align=None, bold=False, size=None, color=None, after=8):
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(after)
        if align is not None: p.alignment = align
        r = p.add_run(text); r.bold = bold; r.font.name = DFONT
        if size: r.font.size = Pt(size)
        r.font.color.rgb = color or ink
        return p

    para(data["name"], align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=18, color=accent, after=1)
    c = data.get("contact", {})
    bits = [c.get("location"), c.get("phone"), c.get("email"), c.get("linkedin")]
    para("  •  ".join(b for b in bits if b), align=WD_ALIGN_PARAGRAPH.CENTER, size=9,
         color=muted, after=10)
    if data.get("date"): para(data["date"], after=6)
    if data.get("company"): para(data["company"], bold=True, after=1)
    if data.get("role"): para("Re: " + data["role"], after=8)
    para(data.get("salutation") or ("Dear %s," % data.get("recipient", "Hiring Team")), after=8)
    for b in data.get("body", []): para(b, after=8)
    para(data.get("closing", "Sincerely,"), after=14)
    para(data["name"])
    doc.save(path)


def main():
    if len(sys.argv) < 2:
        sys.exit("usage: render_cover_letter.py <cover-letter.json> [modern|classic]")
    spec = sys.argv[1]; data = json.load(open(spec))
    style_arg = sys.argv[2] if len(sys.argv) > 2 else None
    style_name, theme = resolve_theme(data, style_arg)
    base = os.path.splitext(os.path.abspath(spec))[0]
    build_docx(data, base + ".docx", theme); print("DOCX: %s" % (base + ".docx"))
    try:
        pages = build_pdf(data, base + ".pdf", theme)
        print("PDF:  %s" % (base + ".pdf"))
        print("STYLE=%s  PAGES=%d  -> %s" % (style_name, pages, "OK" if pages == 1 else "TRIM to one page"))
    except Exception as e:
        print("PDF build failed: %s" % e)


if __name__ == "__main__":
    main()
