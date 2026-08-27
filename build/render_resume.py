#!/usr/bin/env python3
"""
render_resume.py — build a professional, cleanly parsing resume.

Outputs (same dir/basename as the input JSON):
  - <base>.pdf   ReportLab, real selectable text (never image-only), EXACT pagination.
  - <base>.docx  python-docx, native editable Word, for portals that request Word.

Prints a report the caller uses for readability QA — there is no page-count
target; the numbers exist to catch an over-long resume or a near-empty
trailing page:
  STYLE=<name>  PAGES=<n>  LAST_PAGE_FILL=<0..1>  -> <verdict>

Two interchangeable looks (pick per role):
  - "modern"  sans-serif (Helvetica/Calibri) + navy accents — tech/SaaS/growth.
  - "classic" serif (Times / Times New Roman), black, no color — traditional,
              industrial, finance, legal, government.
Both are single column, standard headings, real text (no images): professional
AND clean-parsing. Choose the style via the 2nd CLI arg or a "style" key in JSON.

JSON schema:
{
  "style": "modern|classic",   # optional; default modern (CLI arg overrides)
  "name": "...", "headline": "...(optional honest self-description)...",
  "contact": {"location":"","phone":"","email":"","linkedin":""},
  "summary": "...",
  "experience": [{"title":"","company":"","location":"","dates":"","bullets":["",""]}],
  "competencies": [{"category":"","items":["",""]}],
  "education": ["..."], "certifications": ["..."]
}

Usage:  python3 build/render_resume.py <resume.json> [modern|classic]
"""
import json, os, re, sys

# ---------- chronology enforcement -------------------------------------------
# Recruiters and screeners read work history as a timeline. A resume whose jobs
# are not in reverse-chronological order reads as sloppy or evasive and gets
# called out in interviews. This is enforced here, in the renderer, so that no
# resume can ever be produced out of order regardless of how the JSON was built.

_MONTHS = {m: i for i, m in enumerate(
    ["jan", "feb", "mar", "apr", "may", "jun",
     "jul", "aug", "sep", "oct", "nov", "dec"], start=1)}

# Anything meaning "still there" sorts above every real end date.
_PRESENT = {"present", "current", "now", "today", "ongoing"}

_DATE_SPLIT = re.compile(r"\s*(?:[-–—]|\bto\b)\s*", re.I)


def _parse_point(s):
    """Parse one end of a date range into a sortable (year, month) tuple."""
    s = (s or "").strip().strip(".,")
    if not s:
        return (0, 0)
    if s.lower() in _PRESENT:
        return (9999, 12)
    mon = 0
    m = re.search(r"[A-Za-z]{3,}", s)
    if m:
        mon = _MONTHS.get(m.group(0)[:3].lower(), 0)
    y = re.search(r"(19|20)\d{2}", s)
    return (int(y.group(0)) if y else 0, mon)


def _date_key(job):
    """Sort key for one job: most recent end date first, then most recent start.

    Jobs with unparseable dates keep (0, 0) and fall to the bottom rather than
    scrambling the entries around them.
    """
    parts = _DATE_SPLIT.split(str(job.get("dates", "")).strip(), maxsplit=1)
    start = _parse_point(parts[0] if parts else "")
    end = _parse_point(parts[1]) if len(parts) > 1 else start
    return (end, start)


def enforce_chronology(data):
    """Sort `experience` reverse-chronologically in place.

    Returns a human-readable note when the order actually changed, so the caller
    can report the correction, else None.
    """
    jobs = data.get("experience")
    if not jobs or len(jobs) < 2:
        return None
    before = [id(j) for j in jobs]
    ordered = sorted(jobs, key=_date_key, reverse=True)
    if [id(j) for j in ordered] == before:
        return None
    data["experience"] = ordered
    return " -> ".join("%s (%s)" % (j.get("company", j.get("title", "?")),
                                    j.get("dates", "?")) for j in ordered)

# ---------- shared style constants -------------------------------------------
INK   = (0x1a, 0x1a, 0x1a)
MUTED = (0x55, 0x55, 0x55)
MARGIN_TB_IN = 0.5
MARGIN_LR_IN = 0.6
BODY_PT   = 10.3

# Interchangeable themes. "accent" colors the name, section headings, rule
# lines, company names, and bullets; classic uses ink (no color accent).
THEMES = {
    "modern": {
        "pdf_regular": "Helvetica", "pdf_bold": "Helvetica-Bold", "pdf_italic": "Helvetica-Oblique",
        "docx_font": "Calibri",
        "accent": (0x16, 0x3a, 0x5f), "accent_hex": "163a5f",
    },
    "classic": {
        "pdf_regular": "Times-Roman", "pdf_bold": "Times-Bold", "pdf_italic": "Times-Italic",
        "docx_font": "Times New Roman",
        "accent": INK, "accent_hex": "1a1a1a",
    },
}


def resolve_theme(data, argv_style=None):
    name = (argv_style or data.get("style") or "modern").lower()
    if name not in THEMES:
        name = "modern"
    return name, THEMES[name]


# Map common non-WinAnsi punctuation/symbols to safe equivalents so they don't
# render as missing-glyph boxes in the ReportLab PDF (Helvetica/Times WinAnsi).
# Chars already in CP1252 (en/em dash, curly quotes, bullet, ellipsis, TM) are kept.
_SUBS = {"→": "->", "←": "<-", "⇒": "=>", "↦": "->",
         "−": "-", "‒": "-", "―": "-", "‐": "-", "‑": "-",
         "·": "-", "▪": "-", "●": "-", "◦": "-", "‣": "-",
         " ": " ", " ": " ", " ": " ", " ": " "}


def _san(s):
    """Down-map characters outside WinAnsi (CP1252) coverage to ASCII so
    pasted-JD glyphs (arrows, exotic dashes/bullets, odd spaces) don't render
    as boxes. Anything still uncovered falls back to '?' rather than tofu."""
    s = str(s)
    for k, v in _SUBS.items():
        s = s.replace(k, v)
    return s.encode("cp1252", "replace").decode("cp1252")


def _esc(s):
    """Sanitize to WinAnsi, then escape for ReportLab's mini-XML parser (prevents
    markup injection, entity corruption like 'R&D'->'R&D;', and silent drops of
    <word> patterns)."""
    return _san(s).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


# ============================ PDF (ReportLab) ================================
def build_pdf(data, path, theme):
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.units import inch
    from reportlab.lib.colors import Color
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.platypus import (BaseDocTemplate, PageTemplate, Frame, Paragraph,
                                    Spacer, Table, TableStyle, HRFlowable, KeepTogether)

    F_REG, F_BOLD, F_ITAL = theme["pdf_regular"], theme["pdf_bold"], theme["pdf_italic"]
    accent = Color(*[c / 255 for c in theme["accent"]])
    ink = Color(*[c / 255 for c in INK]); muted = Color(*[c / 255 for c in MUTED])
    ahex = "#" + theme["accent_hex"]
    pw, ph = letter
    ml = MARGIN_LR_IN * inch; mt = MARGIN_TB_IN * inch
    usable = pw - 2 * ml; frameH = ph - 2 * mt

    name_st = ParagraphStyle("name", fontName=F_BOLD, fontSize=21,
                             textColor=accent, alignment=TA_CENTER, spaceAfter=1, leading=24)
    head_st = ParagraphStyle("head", fontName=F_REG, fontSize=11.5,
                             textColor=muted, alignment=TA_CENTER, spaceAfter=2, leading=14)
    contact_st = ParagraphStyle("contact", fontName=F_REG, fontSize=9,
                                textColor=muted, alignment=TA_CENTER, spaceAfter=2, leading=12)
    sec_st = ParagraphStyle("sec", fontName=F_BOLD, fontSize=10.5,
                            textColor=accent, spaceBefore=8, spaceAfter=1, leading=13)
    body_st = ParagraphStyle("body", fontName=F_REG, fontSize=BODY_PT,
                             textColor=ink, leading=BODY_PT + 3, spaceAfter=2, alignment=TA_LEFT)
    role_l = ParagraphStyle("role_l", fontName=F_BOLD, fontSize=10.5,
                            textColor=ink, leading=13)
    role_r = ParagraphStyle("role_r", fontName=F_REG, fontSize=9,
                            textColor=muted, alignment=TA_RIGHT, leading=13)
    loc_st = ParagraphStyle("loc", fontName=F_ITAL, fontSize=9,
                            textColor=muted, leading=11, spaceAfter=1)
    bullet_st = ParagraphStyle("bul", fontName=F_REG, fontSize=BODY_PT, textColor=ink,
                               leading=BODY_PT + 3, leftIndent=11, bulletIndent=0,
                               bulletFontName=F_REG, spaceAfter=1.5)

    def rule():
        return HRFlowable(width="100%", thickness=0.7, color=accent,
                          spaceBefore=1, spaceAfter=3, lineCap="round")

    def section(title):
        return [Paragraph(title.upper(), sec_st), rule()]

    def role_header(job):
        left = Paragraph("<b>%s</b>&nbsp;&nbsp;|&nbsp;&nbsp;<font color='%s'>%s</font>"
                         % (_esc(job["title"]), ahex, _esc(job.get("company", ""))), role_l)
        right = Paragraph(_esc(job.get("dates", "")), role_r)
        t = Table([[left, right]], colWidths=[usable * 0.72, usable * 0.28])
        t.setStyle(TableStyle([("VALIGN", (0, 0), (-1, -1), "TOP"),
                               ("LEFTPADDING", (0, 0), (-1, -1), 0),
                               ("RIGHTPADDING", (0, 0), (-1, -1), 0),
                               ("TOPPADDING", (0, 0), (-1, -1), 0),
                               ("BOTTOMPADDING", (0, 0), (-1, -1), 0)]))
        return t

    def make_story():
        s = [Paragraph(_esc(data["name"]), name_st)]
        if data.get("headline"):
            s.append(Paragraph(_esc(data["headline"]), head_st))
        c = data.get("contact", {})
        bits = [c.get("location"), c.get("phone"), c.get("email"), c.get("linkedin")]
        s.append(Paragraph("&nbsp;&nbsp;&bull;&nbsp;&nbsp;".join(_esc(b) for b in bits if b), contact_st))

        if data.get("summary"):
            s += section("Summary")
            s.append(Paragraph(_esc(data["summary"]), body_st))

        if data.get("experience"):
            s += section("Professional Experience")
            for job in data["experience"]:
                head = [role_header(job)]
                if job.get("location"):
                    head.append(Paragraph(_esc(job["location"]), loc_st))
                bl = job.get("bullets", [])
                first = [Paragraph(_esc(bl[0]), bullet_st, bulletText="•")] if bl else []
                s.append(KeepTogether(head + first))
                for b in bl[1:]:
                    s.append(Paragraph(_esc(b), bullet_st, bulletText="•"))
                s.append(Spacer(1, 3))

        if data.get("competencies"):
            s += section("Core Competencies")
            for g in data["competencies"]:
                s.append(Paragraph("<b>%s:</b>&nbsp;&nbsp;%s"
                                   % (_esc(g["category"]), ", ".join(_esc(i) for i in g["items"])), body_st))

        if data.get("education"):
            s += section("Education")
            for ln in data["education"]:
                s.append(Paragraph(_esc(ln), body_st))
        if data.get("certifications"):
            s += section("Certifications")
            for ln in data["certifications"]:
                s.append(Paragraph(_esc(ln), body_st))
        return s

    frame_top = mt + frameH

    class _Doc(BaseDocTemplate):
        _last_page = 1
        _last_y = mt
        def afterFlowable(self, flowable):
            self._last_page = self.canv.getPageNumber()
            try:
                self._last_y = self.frame._y      # layout cursor after this flowable
            except Exception:
                pass

    doc = _Doc(path, pagesize=letter, leftMargin=ml, rightMargin=ml,
               topMargin=mt, bottomMargin=mt, title="Resume")
    doc.addPageTemplates([PageTemplate(id="main",
                          frames=[Frame(ml, mt, usable, frameH, id="f",
                                        leftPadding=0, rightPadding=0,
                                        topPadding=0, bottomPadding=0)])])
    doc.build(make_story())
    pages = doc._last_page
    used_last = frame_top - doc._last_y          # height consumed on the final page
    last_fill = used_last / frameH
    return pages, max(0.0, min(1.0, last_fill))


# ============================ DOCX (python-docx) =============================
def build_docx(data, path, theme):
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_TAB_ALIGNMENT, WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    DFONT = theme["docx_font"]
    ink = RGBColor(*INK); accent = RGBColor(*theme["accent"]); muted = RGBColor(*MUTED)
    ahex = theme["accent_hex"]

    def run(p, text, *, bold=False, italic=False, size=None, color=None, caps=False):
        r = p.add_run(text); r.bold = bold; r.italic = italic
        if size: r.font.size = Pt(size)
        r.font.color.rgb = color or ink; r.font.name = DFONT
        if caps: r.font.all_caps = True
        return r

    def bottom_border(p):
        pPr = p._p.get_or_add_pPr(); pbdr = OxmlElement("w:pBdr")
        b = OxmlElement("w:bottom")
        b.set(qn("w:val"), "single"); b.set(qn("w:sz"), "6")
        b.set(qn("w:space"), "2"); b.set(qn("w:color"), ahex)
        pbdr.append(b); pPr.append(pbdr)

    def section(title):
        p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(7)
        p.paragraph_format.space_after = Pt(3)
        run(p, title.upper(), bold=True, size=10.5, color=accent, caps=True)
        bottom_border(p)

    doc = Document()
    st = doc.styles["Normal"]; st.font.name = DFONT; st.font.size = Pt(BODY_PT)
    st.font.color.rgb = ink
    rpr = st.element.get_or_add_rPr(); rf = rpr.get_or_add_rFonts()
    for a in ("w:ascii", "w:hAnsi", "w:cs"): rf.set(qn(a), DFONT)
    st.paragraph_format.line_spacing = 1.04
    st.paragraph_format.space_after = Pt(0)
    sec = doc.sections[0]
    sec.top_margin = sec.bottom_margin = Inches(MARGIN_TB_IN)
    sec.left_margin = sec.right_margin = Inches(MARGIN_LR_IN)
    usable = (sec.page_width - sec.left_margin - sec.right_margin) / 914400.0

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(1)
    run(p, data["name"], bold=True, size=21, color=accent)
    if data.get("headline"):
        h = doc.add_paragraph(); h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        h.paragraph_format.space_after = Pt(2); run(h, data["headline"], size=11.5, color=muted)
    c = data.get("contact", {})
    bits = [c.get("location"), c.get("phone"), c.get("email"), c.get("linkedin")]
    cp = doc.add_paragraph(); cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_after = Pt(4)
    run(cp, "  •  ".join(b for b in bits if b), size=9, color=muted)

    if data.get("summary"):
        section("Summary"); p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2)
        run(p, data["summary"])

    if data.get("experience"):
        section("Professional Experience")
        for job in data["experience"]:
            p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.tab_stops.add_tab_stop(Inches(usable), WD_TAB_ALIGNMENT.RIGHT)
            run(p, job["title"], bold=True, size=10.5, color=ink)
            if job.get("company"): run(p, "  |  " + job["company"], size=10.5, color=accent)
            run(p, "\t" + job.get("dates", ""), size=9, color=muted)
            if job.get("location"):
                sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(1)
                run(sp, job["location"], italic=True, size=9, color=muted)
            for b in job.get("bullets", []):
                bp = doc.add_paragraph(); pf = bp.paragraph_format
                pf.left_indent = Inches(0.18); pf.first_line_indent = Inches(-0.18)
                pf.space_after = Pt(1.5)
                run(bp, "•  ", color=accent); run(bp, b)

    if data.get("competencies"):
        section("Core Competencies")
        for g in data["competencies"]:
            p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(1.5)
            run(p, g["category"] + ":  ", bold=True); run(p, ", ".join(g["items"]))

    for title, key in (("Education", "education"), ("Certifications", "certifications")):
        if data.get(key):
            section(title)
            for ln in data[key]:
                p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(1.5); run(p, ln)

    doc.save(path)


def main():
    if len(sys.argv) < 2:
        sys.exit("usage: render_resume.py <resume.json> [modern|classic]")
    spec = sys.argv[1]; data = json.load(open(spec))

    # Never ship an out-of-order work history. Fix the source JSON too, so the
    # spec and the rendered documents can't disagree.
    fixed = enforce_chronology(data)
    if fixed:
        with open(spec, "w") as fh:
            json.dump(data, fh, indent=2, ensure_ascii=False)
            fh.write("\n")
        print("CHRONOLOGY: reordered work history to reverse-chronological: %s" % fixed)
    else:
        print("CHRONOLOGY: OK — reverse-chronological")

    style_arg = sys.argv[2] if len(sys.argv) > 2 else None
    style_name, theme = resolve_theme(data, style_arg)
    base = os.path.splitext(os.path.abspath(spec))[0]
    pdf_path, docx_path = base + ".pdf", base + ".docx"

    build_docx(data, docx_path, theme)
    print("DOCX: %s" % docx_path)
    try:
        pages, fill = build_pdf(data, pdf_path, theme)
        print("PDF:  %s" % pdf_path)
        if pages > 2:
            verdict = ("%d pages — long. Trim to the most relevant, substantiated "
                       "content; do not keep length for its own sake." % pages)
        elif pages > 1 and fill < 0.15:
            verdict = ("near-empty trailing page (fill %.2f) — tighten to %d page(s) "
                       "or let real content carry it. Never pad." % (fill, pages - 1))
        else:
            verdict = "OK — %d page(s), no near-empty trailing page" % pages
        print("STYLE=%s  PAGES=%d  LAST_PAGE_FILL=%.2f  -> %s" % (style_name, pages, fill, verdict))
    except Exception as e:
        print("PDF build failed: %s" % e)
        print("STYLE=%s  PAGES=?  (open the .docx to verify pagination)" % style_name)


if __name__ == "__main__":
    main()
